from docx import Document

# New optimised code
def extract_docx_text(docx_path):
    """Extract text from DOCX"""
    doc = Document(docx_path)
    all_text = []
    # Extract regular text
    full_text = [p.text for p in doc.paragraphs if p.text.strip()]
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(' | '.join(row_text))
    if full_text:
        all_text.append("Document Text:\n" + '\n'.join(full_text))
    return '\n\n'.join(all_text)