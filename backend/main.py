from __future__ import annotations

import os
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional
from uuid import uuid4

from fastapi import FastAPI, File, HTTPException, UploadFile, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, PlainTextResponse, StreamingResponse
from langchain.schema import Document
from pydantic import BaseModel
import uvicorn
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io

from backend.poc_app import create_chunks, create_vector_db, generate_flashcards, generate_notes, generate_quiz, answer_question


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".jpg", ".jpeg", ".png"}


class FileInfo(BaseModel):
    filename: str
    chunkCount: int

class UploadResponse(BaseModel):
    fileId: str
    files: List[FileInfo]
    totalChunks: int


class GenerateRequest(BaseModel):
    fileId: str
    numCards: Optional[int] = None
    difficulty: Optional[str] = None


class FlashcardDTO(BaseModel):
    q: str
    a: str


class GenerateResponse(BaseModel):
    cards: List[FlashcardDTO]


@dataclass
class SessionEntry:
    file_paths: List[Path]
    chunks: List[Document]
    filenames: List[str]
    vector_store: any = None


sessions: Dict[str, SessionEntry] = {}

app = FastAPI(title="Flashcard Generator API")

# CORS configuration - allow all origins for Hugging Face Spaces deployment
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for Hugging Face Spaces
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/health")
def health_check():
    return {"status": "ok"}


@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    """
    Ensure all HTTP errors are returned as plain text instead of JSON.
    """
    detail = exc.detail
    if not isinstance(detail, str):
        detail = str(detail)
    return PlainTextResponse(detail, status_code=exc.status_code)


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception):
    """
    Fallback handler to return plain-text error messages for unexpected errors.
    """
    return PlainTextResponse("An unexpected error occurred while processing your request.", status_code=500)


@app.post("/api/upload", response_model=UploadResponse)
async def upload_files(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="At least one file is required.")
    
    all_chunks = []
    saved_paths = []
    file_infos = []
    filenames = []
    
    for file in files:
        if not file.filename:
            # Clean up already saved files
            for path in saved_paths:
                if path.exists():
                    path.unlink()
            raise HTTPException(status_code=400, detail="All files must have filenames.")
        
        suffix = Path(file.filename).suffix.lower()
        if suffix not in ALLOWED_EXTENSIONS:
            # Clean up already saved files
            for path in saved_paths:
                if path.exists():
                    path.unlink()
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.filename}")
        
        file_bytes = await file.read()
        if not file_bytes:
            # Clean up already saved files
            for path in saved_paths:
                if path.exists():
                    path.unlink()
            raise HTTPException(status_code=400, detail=f"File is empty: {file.filename}")
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            saved_path = Path(tmp.name)
        
        chunks = create_chunks(saved_path)
        if not chunks:
            # Clean up this file and all previously saved files
            if saved_path.exists():
                saved_path.unlink()
            for path in saved_paths:
                if path.exists():
                    path.unlink()
            raise HTTPException(status_code=422, detail=f"Unable to extract text from file: {file.filename}")
        
        saved_paths.append(saved_path)
        all_chunks.extend(chunks)
        filenames.append(file.filename)
        file_infos.append(FileInfo(filename=file.filename, chunkCount=len(chunks)))
    
    file_id = str(uuid4())
    sessions[file_id] = SessionEntry(file_paths=saved_paths, chunks=all_chunks, filenames=filenames)
    
    return UploadResponse(fileId=file_id, files=file_infos, totalChunks=len(all_chunks))


@app.post("/api/generate", response_model=GenerateResponse)
async def generate(file_input: GenerateRequest):
    entry = sessions.get(file_input.fileId)
    if not entry:
        raise HTTPException(status_code=404, detail="Upload session not found.")

    combined_text = " ".join(chunk.page_content for chunk in entry.chunks).strip()
    if not combined_text:
        raise HTTPException(status_code=422, detail="No text found for flashcard generation.")

    num_cards = file_input.numCards if file_input.numCards and file_input.numCards > 0 else 10
    difficulty = file_input.difficulty if file_input.difficulty else "medium"

    cards = generate_flashcards(combined_text, num_cards=num_cards, difficulty=difficulty)
    # Create vector DB with metadata from all files if not already created
    if not entry.vector_store:
        docs_for_db = [Document(page_content=combined_text, metadata={"sources": ", ".join(entry.filenames)})]
        entry.vector_store = create_vector_db(docs_for_db)

    card_payload = [FlashcardDTO(q=card.q, a=card.a) for card in cards]

    return GenerateResponse(cards=card_payload)


class NotesResponse(BaseModel):
    title: str
    summary: str
    keyPoints: List[str]
    detailedNotes: str


@app.post("/api/generate-notes", response_model=NotesResponse)
async def generate_consolidated_notes(file_input: GenerateRequest):
    entry = sessions.get(file_input.fileId)
    if not entry:
        raise HTTPException(status_code=404, detail="Upload session not found.")
    
    combined_text = " ".join(chunk.page_content for chunk in entry.chunks).strip()
    if not combined_text:
        raise HTTPException(status_code=422, detail="No text found for notes generation.")
    
    notes = generate_notes(combined_text)
    
    # Create vector DB with metadata from all files if not already created
    if not entry.vector_store:
        docs_for_db = [Document(page_content=combined_text, metadata={"sources": ", ".join(entry.filenames)})]
        entry.vector_store = create_vector_db(docs_for_db)
    
    return NotesResponse(
        title=notes.title,
        summary=notes.summary,
        keyPoints=notes.key_points,
        detailedNotes=notes.detailed_notes
    )


class QuizQuestionDTO(BaseModel):
    question: str
    options: List[str]
    correctAnswer: str
    explanation: str


class QuizResponse(BaseModel):
    questions: List[QuizQuestionDTO]


class QuizRequest(BaseModel):
    fileId: str
    numQuestions: Optional[int] = None
    difficulty: Optional[str] = None


@app.post("/api/generate-quiz", response_model=QuizResponse)
async def generate_quiz_endpoint(quiz_input: QuizRequest):
    entry = sessions.get(quiz_input.fileId)
    if not entry:
        raise HTTPException(status_code=404, detail="Upload session not found.")
    
    combined_text = " ".join(chunk.page_content for chunk in entry.chunks).strip()
    if not combined_text:
        raise HTTPException(status_code=422, detail="No text found for quiz generation.")
    
    num_questions = quiz_input.numQuestions if quiz_input.numQuestions and quiz_input.numQuestions > 0 else 5
    difficulty = quiz_input.difficulty if quiz_input.difficulty else "medium"
    
    quiz = generate_quiz(combined_text, num_questions=num_questions, difficulty=difficulty)
    
    # Create vector DB if not already created
    if not entry.vector_store:
        docs_for_db = [Document(page_content=combined_text, metadata={"sources": ", ".join(entry.filenames)})]
        entry.vector_store = create_vector_db(docs_for_db)
    
    questions_payload = [
        QuizQuestionDTO(
            question=q.question,
            options=q.options,
            correctAnswer=q.correct_answer,
            explanation=q.explanation
        )
        for q in quiz.questions
    ]
    
    return QuizResponse(questions=questions_payload)


class ChatRequest(BaseModel):
    fileId: str
    question: str


class ChatResponse(BaseModel):
    answer: str
    sources: str


@app.post("/api/chat", response_model=ChatResponse)
async def chat_endpoint(chat_input: ChatRequest):
    entry = sessions.get(chat_input.fileId)
    if not entry:
        raise HTTPException(status_code=404, detail="Upload session not found.")
    
    # Create vector DB if not already created
    if not entry.vector_store:
        combined_text = " ".join(chunk.page_content for chunk in entry.chunks).strip()
        docs_for_db = [Document(page_content=combined_text, metadata={"sources": ", ".join(entry.filenames)})]
        entry.vector_store = create_vector_db(docs_for_db)
    
    result = answer_question(chat_input.question, entry.vector_store)
    
    return ChatResponse(
        answer=result["answer"],
        sources=result["sources"]
    )


@app.get("/api/download-notes/{file_id}")
async def download_notes(file_id: str):
    entry = sessions.get(file_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Upload session not found.")
    
    combined_text = " ".join(chunk.page_content for chunk in entry.chunks).strip()
    if not combined_text:
        raise HTTPException(status_code=422, detail="No text found for notes generation.")
    
    # Generate notes
    notes = generate_notes(combined_text)
    
    # Create DOCX document
    doc = DocxDocument()
    
    # Add title
    title = doc.add_heading(notes.title, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add summary section
    doc.add_heading('Summary', level=1)
    summary_para = doc.add_paragraph(notes.summary)
    summary_para.paragraph_format.space_after = Pt(12)
    
    # Add key points section
    doc.add_heading('Key Points', level=1)
    for point in notes.key_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Add detailed notes section
    doc.add_heading('Detailed Notes', level=1)
    detailed_para = doc.add_paragraph(notes.detailed_notes)
    detailed_para.paragraph_format.space_after = Pt(12)
    
    # Save to BytesIO
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    # Return as streaming response
    filename = f"{notes.title.replace(' ', '_')}.docx"
    return StreamingResponse(
        docx_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# Serve static files (frontend) - must be after API routes
static_dir = Path(__file__).parent.parent / "static"
if static_dir.exists():
    # Mount static assets
    app.mount("/assets", StaticFiles(directory=str(static_dir / "assets")), name="assets")
    
    # Serve index.html for all non-API routes (SPA fallback)
    @app.get("/{full_path:path}")
    async def serve_spa(full_path: str):
        # Don't serve index.html for API routes
        if full_path.startswith("api/"):
            raise HTTPException(status_code=404, detail="Not found")
        
        index_path = static_dir / "index.html"
        if index_path.exists():
            return FileResponse(str(index_path))
        raise HTTPException(status_code=404, detail="Not found")


if __name__ == "__main__":
    port = int(os.getenv("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=False)

